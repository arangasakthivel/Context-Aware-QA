import sys
import os
import re

import numpy as np
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter
from logger import logging
from exception import customexception

from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import pandas as pd
from bs4 import BeautifulSoup

from pdf2image import convert_from_bytes
import pytesseract


# -------------------- TEXT NORMALIZATION --------------------
def normalize_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"([a-z])([A-Z])", r"\1 \2", text)
    return text.strip()


# -------------------- GARBLED TEXT CHECK --------------------
def is_garbled_text(text: str, threshold: float = 0.25) -> bool:
    if not text.strip():
        return True
    non_ascii = sum(1 for c in text if ord(c) > 127)
    return (non_ascii / max(len(text), 1)) > threshold


# -------------------- OCR IMAGE PREPROCESSING --------------------
def preprocess_image_for_ocr(pil_image):
    image = np.array(pil_image)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    # Adaptive thresholding improves scanned PDFs a LOT
    thresh = cv2.adaptiveThreshold(
        gray,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31,
        2,
    )

    # Denoise
    denoised = cv2.medianBlur(thresh, 3)
    return denoised


# -------------------- OCR EXTRACTION --------------------
def extract_text_with_ocr(uploaded_file):
    images = convert_from_bytes(uploaded_file.read(), dpi=300)
    text = ""

    custom_config = r"--oem 3 --psm 6"

    for img in images:
        processed = preprocess_image_for_ocr(img)
        text += pytesseract.image_to_string(
            processed,
            lang="eng",
            config=custom_config
        ) + "\n"

    return normalize_text(text)


# -------------------- PDF EXTRACTION --------------------
def extract_text_from_pdf(uploaded_file):
    try:
        reader = PdfReader(uploaded_file)
        text = ""

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

        # OCR fallback
        if is_garbled_text(text):
            logging.info("PDF appears scanned or garbled → using OCR")
            uploaded_file.seek(0)
            text = extract_text_with_ocr(uploaded_file)

        return normalize_text(text)

    except Exception as e:
        raise Exception(f"Failed to read PDF: {str(e)}")


# -------------------- OTHER FILE TYPES --------------------
def extract_text_from_txt(uploaded_file):
    return uploaded_file.read().decode("utf-8").strip()


def extract_text_from_docx(uploaded_file):
    doc = DocxDocument(uploaded_file)
    return normalize_text("\n".join(p.text for p in doc.paragraphs if p.text))


def extract_text_from_csv(uploaded_file):
    df = pd.read_csv(uploaded_file)
    return df.to_string(index=False)


def extract_text_from_html(uploaded_file):
    soup = BeautifulSoup(uploaded_file.read(), "html.parser")
    return normalize_text(soup.get_text(separator=" "))


# -------------------- MAIN LOADER --------------------
def load_data(uploaded_file):
    try:
        logging.info("Starting data ingestion...")

        if uploaded_file is None:
            raise ValueError("No file uploaded")

        filename = uploaded_file.name.lower()

        if filename.endswith(".pdf"):
            file_content = extract_text_from_pdf(uploaded_file)

        elif filename.endswith(".txt"):
            file_content = extract_text_from_txt(uploaded_file)

        elif filename.endswith(".docx"):
            file_content = extract_text_from_docx(uploaded_file)

        elif filename.endswith(".csv"):
            file_content = extract_text_from_csv(uploaded_file)

        elif filename.endswith(".html") or filename.endswith(".htm"):
            file_content = extract_text_from_html(uploaded_file)

        else:
            raise ValueError("Unsupported file type")

        if not file_content.strip():
            raise ValueError("Extracted text is empty")

        splitter = SentenceSplitter(
            chunk_size=512,
            chunk_overlap=20
        )

        chunks = splitter.split_text(file_content)

        documents = [
            Document(
                text=chunk,
                metadata={
                    "source": filename,
                    "file_type": filename.split(".")[-1]
                }
            )
            for chunk in chunks
        ]

        logging.info("Document successfully processed")
        return documents

    except Exception as e:
        logging.error("Data ingestion failed")
        raise customexception(e, sys)
